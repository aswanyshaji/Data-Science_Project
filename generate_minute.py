import openai
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT

# Set OpenAI API key from environment variable
openai.api_key = os.getenv("OPENAI_API_KEY")

def generate_minutes(transcript_path, meeting_date):
    try:
        with open(transcript_path, "r", encoding="utf-8") as f:
            transcript = f.read()

        results = process_transcription(transcript)

        # Prepare the text minutes content
        minutes_text = (
            f"Date of Meeting: {meeting_date}\n\n"
            f"Abstract Summary:\n{results['abstract_summary']}\n\n"
            f"Key Points:\n" + "\n".join(results['key_points']) + "\n\n"
            f"Action Items:\n" + "\n".join(results['action_items'])
        )

        # Save the minutes as a text file
        minutes_path = os.path.splitext(transcript_path)[0] + "_minutes.txt"
        with open(minutes_path, "w", encoding="utf-8") as f:
            f.write(minutes_text)
        
        # Create the Word document
        doc_path = os.path.join(os.path.dirname(transcript_path), f"{meeting_date}_minutes.docx")
        create_word_doc(results, doc_path, meeting_date)
        
        print("Minutes generation completed.")
        return minutes_path, doc_path
    except Exception as e:
        print(f"Error in generate_minutes: {str(e)}")
        raise e

def process_transcription(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items
    }

def abstract_summary_extraction(transcription):
    prompt = f"""
    You are an expert in summarizing meeting transcripts. Please provide a concise, meaningful, and coherent abstract summary of the text below. 
    Focus on capturing the essence of the discussions, decisions made, and any critical points. Ensure that the summary is well-structured and provides a clear overview of the meeting.

    Text to summarize:
    {transcription}

    Abstract Summary:
    """
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {"role": "system", "content": "You are a skilled summarization AI."},
            {"role": "user", "content": prompt}
        ]
    )
    return response.choices[0].message.content.strip()

def key_points_extraction(transcription):
    prompt = f"""
    You are an expert in extracting key points from meeting transcripts. Identify and list the most important points discussed in the text below. 
    These should be concise, clear, and relevant. Separate each key point using a bullet point. Avoid using numbering.

    Text:
    {transcription}

    Key Points:
    """
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {"role": "system", "content": "You are a skilled summarization AI."},
            {"role": "user", "content": prompt}
        ]
    )
    return [point.strip() for point in response.choices[0].message.content.split('\n') if point.strip()]

def action_item_extraction(transcription):
    prompt = f"""
    You are an expert in extracting actionable items from meeting transcripts. Extract and list all actionable items that were agreed upon during the meeting. 
    Each action item should be clear and specific. Separate each action item with a bullet point. Avoid using numbering.

    Text:
    {transcription}

    Action Items:
    """
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {"role": "system", "content": "You are a skilled summarization AI."},
            {"role": "user", "content": prompt}
        ]
    )
    return [item.strip() for item in response.choices[0].message.content.split('\n') if item.strip()]

def create_word_doc(results, doc_path, meeting_date):
    doc = Document()
    
    # Add the meeting date as a title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Date of Meeting: {meeting_date}")
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph()  # Add a blank line

    # Add sections to the document
    doc.add_heading('Abstract Summary', level=1)
    abstract_paragraph = doc.add_paragraph(results['abstract_summary'])
    abstract_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    doc.add_paragraph()  # Add a blank line

    doc.add_heading('Key Points', level=1)
    for point in results['key_points']:
        point_paragraph = doc.add_paragraph(point)
        point_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph()  # Add a blank line

    doc.add_heading('Action Items', level=1)
    for item in results['action_items']:
        item_paragraph = doc.add_paragraph(item)
        item_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Save the document
    doc.save(doc_path)