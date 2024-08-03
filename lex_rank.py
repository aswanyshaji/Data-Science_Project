# -*- coding: utf-8 -*-
"""
Created on Thu Jul 8 12:54:10 2024

@author: aswany
"""

import nltk
import numpy as np
import networkx as nx
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize, word_tokenize
from sklearn.feature_extraction.text import TfidfVectorizer
from datasets import load_dataset
import os
from docx import Document


nltk.download('punkt')
nltk.download('stopwords')

def preprocess_text(text):
    stop_words = set(stopwords.words('english'))
    sentences = sent_tokenize(text)
    processed_sentences = []

    for sentence in sentences:
        words = word_tokenize(sentence)
        filtered_words = [word for word in words if word.isalnum() and word.lower() not in stop_words]
        processed_sentences.append(' '.join(filtered_words))

    return sentences, processed_sentences

def build_similarity_matrix(sentences):
    vectorizer = TfidfVectorizer().fit_transform(sentences)
    vectors = vectorizer.toarray()
    similarity_matrix = np.dot(vectors, vectors.T)
    return similarity_matrix

def lexrank(text, summary_length=10):
    original_sentences, processed_sentences = preprocess_text(text)
    similarity_matrix = build_similarity_matrix(processed_sentences)

    # Build similarity graph
    nx_graph = nx.from_numpy_array(similarity_matrix)
    scores = nx.pagerank(nx_graph)

    # Rank sentences
    ranked_sentences = sorted(((scores[i], s) for i, s in enumerate(original_sentences)), reverse=True)

    # Adjust summary_length to not exceed the number of sentences in the text
    summary_length = min(summary_length, len(original_sentences))

    # Select top-ranked sentences
    summary = ' '.join([s for _, s in ranked_sentences[:summary_length]])
    return summary

def save_summary_as_word(summary, filename):
    
    doc = Document()
    doc.add_heading('Meeting Summary', 0)
    doc.add_paragraph(summary)
    doc.save(filename)


if __name__ == "__main__":
    
    dataset = load_dataset("huuuyeah/meetingbank")

    
    output_dir = 'summaries'
    os.makedirs(output_dir, exist_ok=True)

    # Summarize the first five transcripts and save them as Word documents
    for index in range(5):
        print(f"Summarizing transcript {index+1}...")
        transcript = dataset['train']['transcript'][index]
        
        # Dynamically adjust summary_length based on the length of the transcript
        sentences = sent_tokenize(transcript)
        summary_length = min(20, len(sentences))
        
        summary = lexrank(transcript, summary_length=summary_length)
        
        # Save summary as a Word document
        file_name = os.path.join(output_dir, f"summary_{index+1}.docx")
        save_summary_as_word(summary, file_name)
        print(f"Summary saved to {file_name}")